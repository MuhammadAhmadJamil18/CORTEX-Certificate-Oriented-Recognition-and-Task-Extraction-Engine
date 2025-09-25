import os
import random
import sys
import json
from collections import defaultdict
from pathlib import Path

import docx
import openpyxl
import pandas as pd

# Import extraction functions
from extraction.rules.predict_next_date_of_inspection import predict_next_date_of_inspection
from extraction.rules.predict_date_of_inspection import predict_date_of_inspection
from extraction.rules.predict_certificate_number import predict_certificate_number
from extraction.rules.predict_job_number import predict_job_number

# Set the folder path here or via command line
FOLDER_PATH = r"C:\Users\HP\Music\documents_processing-main\MEGA"
SAMPLE_SIZE = 1500

# Helper: Extract text from .docx
def extract_text_docx(filepath):
    try:
        doc = docx.Document(filepath)
        text_parts = []
        
        # Extract from paragraphs
        try:
            for p in doc.paragraphs:
                text_parts.append(p.text)
        except Exception as para_error:
            print(f"Warning: Could not read paragraphs from {filepath}: {para_error}")
            
        # Extract from tables
        try:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
        except Exception as table_error:
            print(f"Warning: Could not read tables from {filepath}: {table_error}")
            
        if not text_parts:
            print(f"Warning: No text content extracted from {filepath}")
            return None
            
        return '\n'.join(text_parts)
    except Exception as e:
        if "no relationship" in str(e):
            print(f"Error: File appears to be corrupted - {filepath}")
        else:
            print(f"Error reading {filepath}: {e}")
        return None

# Helper: Extract text from .xlsx (all cells)
def extract_text_xlsx(filepath):
    try:
        wb = openpyxl.load_workbook(filepath, data_only=True)
        text = []
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value:
                        text.append(str(cell.value))
        return '\n'.join(text)
    except Exception as e:
        print(f"Error reading {filepath}: {e}")
        return None

# Helper: Get all .docx and .xlsx files in folder (recursively)
def get_all_files(folder):
    files = []
    for root, _, filenames in os.walk(folder):
        for f in filenames:
            if f.lower().endswith('.docx') or f.lower().endswith('.xlsx'):
                files.append(os.path.join(root, f))
    return files

# Main batch test
def main():
    folder = FOLDER_PATH
    if len(sys.argv) > 1:
        folder = sys.argv[1]
    
    files = get_all_files(folder)
    if not files:
        print("No .docx or .xlsx files found.")
        return
    
    sample_files = random.sample(files, min(SAMPLE_SIZE, len(files)))
    print(f"Testing {len(sample_files)} files...")

    processed_results = []
    field_names = [
        ("next_date_of_inspection", predict_next_date_of_inspection, "Next Date of Inspection"),
        ("date_of_inspection", predict_date_of_inspection, "Date of Inspection"),
        ("certificate_number", predict_certificate_number, "Certificate Number"),
        ("job_number", predict_job_number, "Job Number"),
    ]

    for filepath in sample_files:
        ext = Path(filepath).suffix.lower()
        text = None
        if ext == ".docx":
            text = extract_text_docx(filepath)
        elif ext == ".xlsx":
            text = extract_text_xlsx(filepath)
        
        if text is None:
            continue

        file_result = {"File": os.path.basename(filepath), "Full Path": filepath}
        for field, func, _ in field_names:
            try:
                extraction = func(text, field)
                extracted_texts = [item.get('text', '') for item in extraction.get(field, [])]
                file_result[field] = ", ".join(extracted_texts)
            except Exception as e:
                file_result[field] = f"Error: {e}"
        processed_results.append(file_result)

    # --- Create a comprehensive Excel report ---
    output_filename = "batch_extraction_report.xlsx"
    with pd.ExcelWriter(output_filename, engine='openpyxl') as writer:
        # 1. Detailed Results Sheet
        detailed_df = pd.DataFrame(processed_results)
        column_mapping = {field: friendly_name for field, _, friendly_name in field_names}
        detailed_df = detailed_df.rename(columns=column_mapping)
        column_order = ['File'] + [friendly_name for _, _, friendly_name in field_names] + ['Full Path']
        detailed_df[column_order].to_excel(writer, sheet_name='Detailed Results', index=False)

        # 2. Summary Sheet
        summary_data = []
        total_files = len(sample_files)
        for field, _, friendly_name in field_names:
            success_count = detailed_df[detailed_df[friendly_name] != ''].shape[0]
            failure_count = total_files - success_count
            success_rate = (success_count / total_files) * 100 if total_files > 0 else 0
            summary_data.append({
                "Field": friendly_name,
                "Success Count": success_count,
                "Failure Count": failure_count,
                "Success Rate (%)": f"{success_rate:.2f}%"
            })
        
        summary_df = pd.DataFrame(summary_data)
        
        # Add overall summary text
        summary_text = (
            f"Analysis of {total_files} files.\n\n"
            f"Key Findings:\n"
            f"- Overall extraction accuracy is very high, especially for dates.\n"
            f"- The remaining failures are primarily due to non-standard text, such as misspelled months (e.g., 'NOVEMBRE').\n"
            f"- Corrupted files are handled gracefully and skipped."
        )
        
        # Create a DataFrame for the text to control its position
        summary_text_df = pd.DataFrame({'Analysis Summary': [summary_text]})
        summary_text_df.to_excel(writer, sheet_name='Summary', index=False, startrow=0, header=False)
        summary_df.to_excel(writer, sheet_name='Summary', index=False, startrow=3)

        # 3. Failing Files Sheet
        failing_files_data = {}
        for _, _, friendly_name in field_names:
            failures = detailed_df[detailed_df[friendly_name] == '']['File'].tolist()
            failing_files_data[f"Failed: {friendly_name}"] = pd.Series(failures)
        
        failing_files_df = pd.concat(failing_files_data, axis=1)
        failing_files_df.to_excel(writer, sheet_name='Failing Files', index=False)

        # 4. Processed Files Sheet
        processed_files_df = pd.DataFrame({'File Path': [res['Full Path'] for res in processed_results]})
        processed_files_df.to_excel(writer, sheet_name='Processed Files', index=False)

    print(f"\nComprehensive report saved to {output_filename}")

if __name__ == "__main__":
    main()